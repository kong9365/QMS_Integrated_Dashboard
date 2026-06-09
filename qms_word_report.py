# -*- coding: utf-8 -*-
"""qms_word_report.py — OOS 경향분석보고서 Word(.docx) 생성기 (KD-MoaQ).

설계 원칙
- **신규 순수함수만 추가**. 기존 표시 함수 ``qms_oos_dashboard_panels.render_oos_report``
  및 도메인/건수기여도 계산 로직은 일절 수정하지 않는다.
- render_oos_report 의 집계(건수기여도 기반 groupby/sum)를 **그대로 미러링**해
  python-docx 표로 출력한다 → 화면 숫자와 보고서 숫자 일치.
- 차트(Plotly)는 kaleido 미설치 환경을 고려해 **데이터 표**로 대체(Word 문서에는
  정확한 수치 표가 차트보다 유용). 추후 kaleido 도입 시 이미지 임베드 확장 가능.

공개 API
- build_oos_trend_report_docx(filtered, safe_pct, completed_keywords, *,
      as_of=None, project_label="OOS (Out of Specification)", filter_note="") -> bytes
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

_MONTH_LABELS = [f"{m}월" for m in range(1, 13)]

# ── 디자인 토큰 (광동 CI) ──
_KD_RED = RGBColor(0xE8, 0x30, 0x08)   # 강조/경고
_KD_RED_HEX = "E83008"
_NAVY = RGBColor(0x1F, 0x3A, 0x5F)     # 헤더/주계열
_NAVY_HEX = "1F3A5F"
_HDR_BG = _NAVY_HEX                      # 하위호환
_GREEN_HEX = "1F9D63"                    # 완료(성공)
_INK = RGBColor(0x1B, 0x23, 0x30)       # 본문 잉크
_SUB = RGBColor(0x51, 0x59, 0x6A)       # 보조 라벨
_MUTED = RGBColor(0xC2, 0xC9, 0xD6)     # 0값 톤다운
_ZEBRA = "F4F7FA"                        # 짝수행 줄무늬
_LINE = "E1E6EF"                         # 옅은 가로선
_TOTAL_BG = "EAF0F8"                     # 합계행 배경
_GREY = RGBColor(0x80, 0x80, 0x80)
_MONO = "Consolas"                       # 숫자/관리번호 등폭

_NUM_NAME = ("건수", "합계", "순위", "비율", "값", "율", "%")
_NUM_RE = re.compile(r"^\s*[-+]?[\d,]+(?:\.\d+)?\s*(?:건|%|일|개|회|점)?\s*$")


def _set_cell_bg(cell, hex_color: str) -> None:
    """표 셀 배경 음영."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_margins(cell, *, top=40, bottom=40, left=90, right=90) -> None:
    """셀 안쪽 여백(twips) — 답답함 제거."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, w in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tc_pr.append(tcMar)


def _set_cell_borders(cell, *, top=None, bottom=None) -> None:
    """셀 가로 테두리만(세로선 없음). 인자 = (hex, sz) 또는 None."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in (("top", top), ("bottom", bottom)):
        if spec is None:
            continue
        color, sz = spec
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")


def _repeat_header(row) -> None:
    """긴 표가 페이지를 넘어가도 머리행 반복(w:tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _is_num_col(name, series) -> bool:
    """숫자(우측정렬·등폭) 컬럼인지 — 이름 또는 값 패턴으로 판정."""
    n = str(name)
    if n in _MONTH_LABELS or any(k in n for k in _NUM_NAME):
        return True
    vals = [str(v).strip() for v in series if str(v).strip()]
    return bool(vals) and all(_NUM_RE.match(v) for v in vals)


def _fmt_num(txt) -> str:
    """정수부 천단위 콤마(접미사 건/%/일 보존)."""
    m = re.match(r"^\s*([-+]?\d+)(\.\d+)?\s*(\D*)\s*$", str(txt))
    if not m:
        return str(txt)
    try:
        head = f"{int(m.group(1)):,}"
    except Exception:
        return str(txt)
    return f"{head}{m.group(2) or ''}{m.group(3) or ''}"


def _is_zero(txt) -> bool:
    return bool(re.match(r"^\s*0\s*(?:건|%|일|개|회|점)?\s*$", str(txt)))


def _style_runs(cell, *, bold=False, white=False, size=10, mono=False, color=None) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            if mono:
                r.font.name = _MONO
            if white:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif color is not None:
                r.font.color.rgb = color


def _add_df_table(doc, df: pd.DataFrame, *, total_label: Optional[str] = None) -> None:
    """DataFrame → Word 표: 가로선만(세로선 없음)·헤더 네이비+KD Red 밑줄·셀 수직중앙·
    숫자열 우측정렬+등폭+콤마·짝수행 줄무늬·합계행 강조·머리행 반복·본문 10pt·0값 회색.
    인덱스 무시(필요 시 reset_index 후 전달)."""
    cols = [str(c) for c in df.columns]
    ncol = len(cols)
    num = [_is_num_col(cols[j], df[df.columns[j]]) for j in range(ncol)]
    table = doc.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # 헤더
    hrow = table.rows[0]
    _repeat_header(hrow)
    for j, c in enumerate(cols):
        cell = hrow.cells[j]
        cell.text = c
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, _NAVY_HEX)
        _set_cell_margins(cell)
        _set_cell_borders(cell, bottom=(_KD_RED_HEX, 12))   # 헤더 아래 KD Red 강조선
        _style_runs(cell, bold=True, white=True, size=10)
    # 데이터
    for ri, (_, row) in enumerate(df.iterrows()):
        is_total = total_label is not None and str(row[df.columns[0]]).strip() == total_label
        cells = table.add_row().cells
        for j, c in enumerate(df.columns):
            cell = cells[j]
            val = row[c]
            txt = "" if pd.isna(val) else str(val)
            if num[j] and txt:
                txt = _fmt_num(txt)
            cell.text = txt
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if num[j] else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_cell_margins(cell)
            _set_cell_borders(cell, bottom=(_LINE, 4))
            _col = _MUTED if (num[j] and _is_zero(txt)) else _INK
            _style_runs(cell, size=10, mono=num[j], bold=is_total, color=_col)
            if is_total:
                _set_cell_bg(cell, _TOTAL_BG)
                _set_cell_borders(cell, top=(_NAVY_HEX, 12))   # 합계행 상단 굵은 네이비선
            elif ri % 2 == 1:
                _set_cell_bg(cell, _ZEBRA)


def _add_kpi_cards(doc, cards) -> None:
    """요약을 KPI 카드 한 줄(큰 숫자+라벨, 상단 색 보더)로. cards=[(label, value, accent_hex), ...]."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, (label, value, accent) in enumerate(cards):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        _set_cell_borders(cell, top=(accent, 28), bottom=(_LINE, 4))
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(value)
        r1.font.size = Pt(20)
        r1.font.bold = True
        r1.font.color.rgb = _NAVY
        r1.font.name = _MONO
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = _SUB


def _heading(doc, text: str, level: int = 2) -> None:
    """색 있는 섹션 제목(KD Red 막대 + 네이비 굵게). ◆/◇ 기호는 막대로 대체."""
    label = text.lstrip("◆◇ ").strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    bar = p.add_run("▌ ")
    bar.font.color.rgb = _KD_RED
    bar.font.size = Pt(13)
    bar.font.bold = True
    t = p.add_run(label)
    t.font.color.rgb = _NAVY
    t.font.size = Pt(13)
    t.font.bold = True


def _completed_mask(filtered: pd.DataFrame, completed_keywords: Tuple[str, ...]) -> pd.Series:
    """render_oos_report 와 동일한 완료 판정(진행상태 contains → 완료여부=='C' → False)."""
    if "진행상태" in filtered.columns:
        return filtered["진행상태"].str.contains("|".join(completed_keywords), case=False, na=False)
    if "완료여부" in filtered.columns:
        return filtered["완료여부"] == "C"
    return pd.Series(False, index=filtered.index)


def build_oos_trend_report_docx(
    filtered: pd.DataFrame,
    safe_pct,
    completed_keywords: Tuple[str, ...],
    *,
    as_of: Optional[str] = None,
    project_label: str = "OOS (Out of Specification)",
    filter_note: str = "",
) -> bytes:
    """OOS 경향분석보고서를 .docx 바이트로 생성(render_oos_report 집계 미러링)."""
    doc = Document()
    # 기본 폰트(맑은 고딕)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "맑은 고딕"
        normal.font.size = Pt(10)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    except Exception:
        pass

    # ── 표지/헤더 ──
    brand = doc.add_paragraph()
    br = brand.add_run("광동제약 품질부문 · KD-MoaQ")
    br.font.size = Pt(10)
    br.font.bold = True
    br.font.color.rgb = _KD_RED

    doc.add_heading("OOS 경향분석 보고서", level=0)

    as_of = as_of or datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.add_run(f"대상: {project_label}    기준일: {as_of}").font.size = Pt(9)
    if filter_note:
        fp = doc.add_paragraph()
        fp.add_run(f"필터: {filter_note}").font.size = Pt(9)

    # 데이터 없음 가드
    if filtered is None or getattr(filtered, "empty", True) or "건수기여도" not in filtered.columns:
        doc.add_paragraph("표시할 OOS 데이터가 없습니다 (필터 결과 0건 또는 건수기여도 컬럼 없음).")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── ◆ 경향 요약 ──
    total_w = float(filtered["건수기여도"].sum())
    comp_w = float(filtered.loc[_completed_mask(filtered, completed_keywords), "건수기여도"].sum())
    comp_rate = safe_pct(comp_w, total_w)
    _heading(doc, "◆ 경향 요약")
    _add_kpi_cards(doc, [
        ("전체 건수", f"{total_w:,.0f}", _NAVY_HEX),
        ("완료 건수", f"{comp_w:,.0f}", _GREEN_HEX),
        ("완료율", f"{comp_rate:.1f}%", _NAVY_HEX),
        ("미완료 건수", f"{total_w - comp_w:,.0f}", _KD_RED_HEX),
    ])

    # ── ◇ 시험종류별 월별 건수 ──
    if "시험종류" in filtered.columns and "월" in filtered.columns:
        pivot = filtered.groupby(["시험종류", "월"])["건수기여도"].sum().round().unstack(fill_value=0)
        for m in range(1, 13):
            if m not in pivot.columns:
                pivot[m] = 0
        pivot = pivot[sorted(pivot.columns)]
        pivot.columns = _MONTH_LABELS
        pivot = pivot.astype(int)
        # [빈 구간 축약] 합계 0인 달 열 제거(합계 열은 항상 유지) + 주석
        nonzero = [m for m in _MONTH_LABELS if int(pivot[m].sum()) > 0]
        dropped = [m for m in _MONTH_LABELS if m not in nonzero]
        pivot = pivot[nonzero] if nonzero else pivot
        pivot["합계"] = pivot[nonzero].sum(axis=1).astype(int) if nonzero else 0
        pivot = pivot.sort_values("합계", ascending=False)
        total_row = pivot.sum().to_frame().T
        total_row.index = ["합계"]
        full = pd.concat([pivot, total_row])
        full.index.name = "시험종류"
        _heading(doc, "◇ 시험종류별 월별 건수")
        _add_df_table(doc, full.reset_index(), total_label="합계")
        if dropped:
            note = doc.add_paragraph()
            nr = note.add_run(f"※ {', '.join(dropped)}: 발생 없음")
            nr.font.size = Pt(8)
            nr.font.color.rgb = _GREY

    # ── ◇ 상위 OOS 시험항목 (Top 10) ──
    if "시험항목" in filtered.columns:
        item = (
            filtered[filtered["시험항목"].notna() & (filtered["시험항목"] != "")]
            .groupby("시험항목")["건수기여도"].sum().round().nlargest(10).reset_index()
        )
        if not item.empty:
            item.columns = ["시험항목", "건수"]
            item.insert(0, "순위", range(1, len(item) + 1))
            item["건수"] = item["건수"].astype(int)
            _heading(doc, "◇ 상위 OOS 시험항목 (Top 10)")
            _add_df_table(doc, item)

    # ── ◇ 제품별 OOS 건수 (Top 15) ──
    if "품목명" in filtered.columns:
        prod = (
            filtered[filtered["품목명"].notna() & (filtered["품목명"] != "")]
            .groupby("품목명")["건수기여도"].sum().round().sort_values(ascending=False).head(15).reset_index()
        )
        if not prod.empty:
            prod.columns = ["품목명", "건수"]
            prod["건수"] = prod["건수"].astype(int)
            _heading(doc, "◇ 제품별 OOS 건수 (Top 15)")
            _add_df_table(doc, prod)

    # ── ◇ 조치현황 / CAPA 현황 ──
    if "CAPA/Action item 필요여부" in filtered.columns:
        has = filtered["CAPA/Action item 필요여부"].fillna("").str.len() > 0
        no_act = filtered["CAPA/Action item 필요여부"].fillna("").str.contains("No Action", case=False, na=False)
        capa_df = filtered[has & ~no_act]
        _heading(doc, "◇ 조치현황 / CAPA 현황")
        if not capa_df.empty:
            vc = capa_df["CAPA/Action item 필요여부"].value_counts()
            _add_df_table(doc, pd.DataFrame({"CAPA 유형": vc.index, "건수": vc.values.astype(int)}))
        else:
            doc.add_paragraph("CAPA 데이터가 없습니다.")

    # ── ◆ 이상발생 원인 분류 분석 (건수 + 비율) ──
    if "이상발생 원인" in filtered.columns:
        cause = (
            filtered[filtered["이상발생 원인"].notna() & (filtered["이상발생 원인"] != "")]
            .groupby("이상발생 원인")["건수기여도"].sum().round().sort_values(ascending=False).reset_index()
        )
        if not cause.empty:
            cause.columns = ["이상발생 원인", "건수"]
            tot = cause["건수"].sum()
            cause["비율 (%)"] = cause["건수"].apply(lambda x: f"{safe_pct(x, tot):.1f}%")
            cause["건수"] = cause["건수"].astype(int)
            _heading(doc, "◆ 이상발생 원인 분류 분석")
            _add_df_table(doc, cause)

    # ── ◆ 확인된 이벤트 분류별 ──
    if "확인된 이벤트 분류" in filtered.columns:
        cls = (
            filtered[filtered["확인된 이벤트 분류"].notna() & (filtered["확인된 이벤트 분류"] != "")]
            .groupby("확인된 이벤트 분류")["건수기여도"].sum().round().sort_values(ascending=False).reset_index()
        )
        if not cls.empty:
            cls.columns = ["이벤트 분류", "건수"]
            cls["건수"] = cls["건수"].astype(int)
            _heading(doc, "◆ 확인된 이벤트 분류별")
            _add_df_table(doc, cls)

    # ── 푸터 ──
    foot = doc.add_paragraph()
    fr = foot.add_run(f"생성: KD-MoaQ · 광동제약 품질부문 · {datetime.now():%Y-%m-%d %H:%M}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = _GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
